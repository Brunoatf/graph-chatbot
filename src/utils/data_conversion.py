import re
import pandas as pd
from io import BytesIO, StringIO
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def extract_table(message: str) -> str:

    '''Extracts a markdown table from a string.
    Returns the table as a string or an empty string if no table is found.'''

    regex_pattern = r"(\|.*\|\n)+"

    table_match = re.search(regex_pattern, message, re.MULTILINE)

    if table_match:
        markdown_table = table_match.group()
        print("Markdown table found:", markdown_table)
    else:
        markdown_table = ""

    return markdown_table

def convert_df_to_excel(df: pd.DataFrame) -> bytes:

    '''Converts a dataframe to an Excel file and returns the file as a byte string.'''

    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False)
    return output.getvalue()

def markdown_table_to_excel(md_table: str) -> bytes:

    '''Converts a markdown table to an Excel file and returns the file as a byte string.'''

    lines = md_table.strip().split('\n')

    lines.pop(1)

    clean_table = '\n'.join([line[1:-1] for line in lines])

    df = pd.read_csv(StringIO(clean_table), sep='|', skipinitialspace=True)
    print("Markdown to dataframe completed:", df.head())

    return convert_df_to_excel(df)

def chat_to_word(messages: dict, user_name: str) -> bytes:

    '''Converts a chat to a Word document and returns the document as a byte string.'''

    document = Document()

    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'

    section = document.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    run = header_paragraph.add_run()
    image_path = '../images/neuralmind.png'
    run.add_picture(image_path, width=Inches(0.5))
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    document.add_heading("Histórico de conversa com o protótipo do NeuralSearchX-Chatbot", 0)
    document.add_heading(f"Data e hora da conversa: {datetime.now().strftime('%d/%m/%Y %H:%M')}", level=1)
    document.add_heading(f"Nome do usuário: {user_name}", level=1)
    document.add_heading("Mensagens:\n", level=1)

    for message in messages:
        p = document.add_paragraph()
        if message["role"] == "assistant":
            role = p.add_run("Assistente: ")
            role.bold = True
            p.add_run(f"{message['content']}")
        elif message["role"] == "user":
            role = p.add_run("Usuário: ")
            role.bold = True
            p.add_run(f"{message['content']}")
    
    arquivo_word_io = BytesIO()
    document.save(arquivo_word_io)
    arquivo_word_io.seek(0)
    return arquivo_word_io